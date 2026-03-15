from docx import Document
import tempfile

class DOCXExporter:
    def export_project(self, project, chapters):
        document = Document()

        document.add_heading(project.title, level=0)

        for chapter in chapters:
            document.add_heading(chapter.title, level=1)
            content = chapter.content or ""
            document.add_paragraph(content)

        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp_file.name)

        return tmp_file.name
